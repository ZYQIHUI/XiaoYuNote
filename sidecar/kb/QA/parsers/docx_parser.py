"""Word (docx) 解析器 — python-docx 段落+表格（r1 6.4 节）"""

from __future__ import annotations

import logging
from pathlib import Path

from ...config import Config
from ...models import Document
from ...utils import extract_date_from_path, desensitize, is_sensitive_file

logger = logging.getLogger("kb.parsers.docx")


def parse_docx(file_path: Path, cfg: Config) -> list[Document]:
    """解析 docx 文件：正文段落按标题样式分段，表格按行级切块。"""
    docs: list[Document] = []
    rel_path = file_path.relative_to(cfg.root_dir).as_posix()
    date, date_source = extract_date_from_path(file_path)

    try:
        from docx import Document as DocxDoc
        doc = DocxDoc(str(file_path))
    except Exception as e:
        logger.error(f"docx 解析失败: {file_path} — {e}")
        return []

    # 收集所有文本块（段落 + 表格）
    all_texts: list[tuple[str, Optional[str]]] = []  # (text, heading)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        heading = None
        if "Heading" in style_name or "标题" in style_name:
            heading = f"{style_name}: {text[:50]}"
        all_texts.append((text, heading))

    # 表格按行级切块
    for table_idx, table in enumerate(doc.tables):
        headers: list[str] = []
        first_row = True
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not any(c for c in cells):
                continue
            if first_row and cells:
                headers = cells
                first_row = False
                continue
            if headers:
                line = "；".join(f"{h}={v}" for h, v in zip(headers, cells) if v)
                all_texts.append((line, None))

    # 凭据脱敏
    full_text = "\n".join(t for t, _ in all_texts)
    clean_text, desens_count = desensitize(full_text)
    sensitive = is_sensitive_file(desens_count, rel_path)
    if sensitive and cfg.security.skip_sensitive_files:
        logger.warning(f"敏感文件跳过: {rel_path}")
        return []

    # 按文本块生成 Document
    for text, heading in all_texts:
        if len(text) < cfg.chunk.min_size:
            continue
        docs.append(Document(
            path=str(file_path), rel_path=rel_path, file_type="docx",
            date=date, date_source=date_source,
            heading=heading, text=text, sensitive=sensitive,
        ))

    return docs if docs else [Document(
        path=str(file_path), rel_path=rel_path, file_type="docx",
        date=date, date_source=date_source,
        text=full_text, sensitive=sensitive,
    )]
